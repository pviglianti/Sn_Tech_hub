"""Tests for assessment report export (Excel and Word).

Verifies:
1. Excel export generates valid .xlsx bytes with expected sheets
2. Word export generates valid .docx bytes with expected sections
3. Export API endpoints return correct content types and filenames
4. Export handles empty assessments gracefully
5. Export includes features, recommendations, and customized artifacts
"""

import io
import json
import pytest
from unittest.mock import patch

from openpyxl import load_workbook
from docx import Document

from src.models import (
    Assessment,
    AssessmentState,
    AssessmentType,
    Feature,
    FeatureScanResult,
    GeneralRecommendation,
    Instance,
    OriginType,
    PipelineStage,
    ReviewStatus,
    Scan,
    ScanResult,
    ScanStatus,
    ScanType,
    Severity,
)
from src.services.report_export import generate_excel_report, generate_word_report


# ---------------------------------------------------------------------------
# Seed helpers
# ---------------------------------------------------------------------------


def _seed_full_assessment(db_session):
    """Create an assessment with scans, results, features, and recommendations."""
    inst = Instance(
        name="export-test-inst",
        url="https://export-test.service-now.com",
        username="admin",
        password_encrypted="encrypted",
    )
    db_session.add(inst)
    db_session.flush()

    asmt = Assessment(
        instance_id=inst.id,
        name="Export Test Assessment",
        number="ASMT0099902",
        assessment_type=AssessmentType.global_app,
        state=AssessmentState.in_progress,
        pipeline_stage=PipelineStage.complete.value,
    )
    db_session.add(asmt)
    db_session.flush()

    scan = Scan(
        assessment_id=asmt.id,
        scan_type=ScanType.metadata,
        name="export-test-scan",
        status=ScanStatus.completed,
    )
    db_session.add(scan)
    db_session.flush()

    # Add customized results
    sr1 = ScanResult(
        scan_id=scan.id,
        sys_id="sys_export_1",
        table_name="sys_script_include",
        name="ExportScript1",
        origin_type=OriginType.modified_ootb,
        review_status=ReviewStatus.reviewed,
        finding_description="A modified script include",
    )
    sr2 = ScanResult(
        scan_id=scan.id,
        sys_id="sys_export_2",
        table_name="sys_script",
        name="ExportBR1",
        origin_type=OriginType.net_new_customer,
        review_status=ReviewStatus.pending_review,
    )
    # OOTB result (should not appear in customized sheet)
    sr3 = ScanResult(
        scan_id=scan.id,
        sys_id="sys_export_ootb",
        table_name="sys_script_include",
        name="OotbScript",
        origin_type=OriginType.ootb_untouched,
    )
    db_session.add_all([sr1, sr2, sr3])
    db_session.flush()

    # Add a feature
    feat = Feature(
        assessment_id=asmt.id,
        name="Export Feature Group",
        description="A test feature for export",
        ai_summary="AI-generated summary for export testing",
        recommendation="Recommend upgrading this feature",
    )
    db_session.add(feat)
    db_session.flush()

    # Link sr1 to the feature
    fsr = FeatureScanResult(feature_id=feat.id, scan_result_id=sr1.id)
    db_session.add(fsr)

    # Add a general recommendation
    rec = GeneralRecommendation(
        assessment_id=asmt.id,
        title="Upgrade Script Includes",
        category="technical_findings",
        severity=Severity.high,
        created_by="test",
        description="Multiple script includes need upgrading.",
    )
    db_session.add(rec)
    db_session.commit()

    return asmt


def _seed_empty_assessment(db_session):
    """Create an assessment with no scan results."""
    inst = Instance(
        name="empty-export-inst",
        url="https://empty-export.service-now.com",
        username="admin",
        password_encrypted="encrypted",
    )
    db_session.add(inst)
    db_session.flush()

    asmt = Assessment(
        instance_id=inst.id,
        name="Empty Export Assessment",
        number="ASMT0099903",
        assessment_type=AssessmentType.global_app,
        state=AssessmentState.in_progress,
    )
    db_session.add(asmt)
    db_session.commit()
    return asmt


# ---------------------------------------------------------------------------
# Excel export tests
# ---------------------------------------------------------------------------


def test_excel_export_returns_valid_xlsx(db_session):
    """generate_excel_report should return valid xlsx bytes."""
    asmt = _seed_full_assessment(db_session)
    data = generate_excel_report(db_session, asmt.id)

    assert isinstance(data, bytes)
    assert len(data) > 100

    # Verify it's a valid workbook
    wb = load_workbook(io.BytesIO(data))
    assert "Summary" in wb.sheetnames
    assert "Customized Artifacts" in wb.sheetnames
    assert "Features" in wb.sheetnames
    assert "Recommendations" in wb.sheetnames


def test_excel_export_summary_sheet_content(db_session):
    """Summary sheet should contain assessment name and statistics."""
    asmt = _seed_full_assessment(db_session)
    data = generate_excel_report(db_session, asmt.id)
    wb = load_workbook(io.BytesIO(data))
    ws = wb["Summary"]

    # Check title row
    assert "Export Test Assessment" in str(ws.cell(row=2, column=1).value)

    # Check stats exist
    cell_values = [ws.cell(row=r, column=1).value for r in range(7, 20)]
    metric_names = [v for v in cell_values if v]
    assert "Total Artifacts" in metric_names
    assert "Customized Artifacts" in metric_names


def test_excel_export_customized_artifacts_count(db_session):
    """Customized Artifacts sheet should contain only customized results."""
    asmt = _seed_full_assessment(db_session)
    data = generate_excel_report(db_session, asmt.id)
    wb = load_workbook(io.BytesIO(data))
    ws = wb["Customized Artifacts"]

    # Row 1 is header; customized count should be 2 (sr1, sr2)
    row_count = 0
    for row in ws.iter_rows(min_row=2, max_col=1, values_only=True):
        if row[0]:
            row_count += 1
    assert row_count == 2


def test_excel_export_features_sheet(db_session):
    """Features sheet should contain feature groups."""
    asmt = _seed_full_assessment(db_session)
    data = generate_excel_report(db_session, asmt.id)
    wb = load_workbook(io.BytesIO(data))
    ws = wb["Features"]

    assert ws.cell(row=2, column=1).value == "Export Feature Group"
    assert ws.cell(row=2, column=4).value == 1  # member count


def test_excel_export_empty_assessment(db_session):
    """Excel export of empty assessment should succeed with zero data rows."""
    asmt = _seed_empty_assessment(db_session)
    data = generate_excel_report(db_session, asmt.id)

    wb = load_workbook(io.BytesIO(data))
    assert "Summary" in wb.sheetnames
    ws = wb["Customized Artifacts"]
    # Only header row, no data rows
    rows = list(ws.iter_rows(min_row=2, max_col=1, values_only=True))
    non_empty = [r for r in rows if r[0]]
    assert len(non_empty) == 0


# ---------------------------------------------------------------------------
# Word export tests
# ---------------------------------------------------------------------------


def test_word_export_returns_valid_docx(db_session):
    """generate_word_report should return valid docx bytes."""
    asmt = _seed_full_assessment(db_session)
    data = generate_word_report(db_session, asmt.id)

    assert isinstance(data, bytes)
    assert len(data) > 100

    # Verify it's a valid Word document
    doc = Document(io.BytesIO(data))
    assert len(doc.paragraphs) > 0


def test_word_export_contains_assessment_info(db_session):
    """Word export should contain assessment name and number."""
    asmt = _seed_full_assessment(db_session)
    data = generate_word_report(db_session, asmt.id)
    doc = Document(io.BytesIO(data))

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Export Test Assessment" in all_text
    assert "ASMT0099902" in all_text


def test_word_export_contains_sections(db_session):
    """Word export should contain expected section headings."""
    asmt = _seed_full_assessment(db_session)
    data = generate_word_report(db_session, asmt.id)
    doc = Document(io.BytesIO(data))

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    heading_texts = " ".join(headings)
    assert "Executive Summary" in heading_texts
    assert "Customized Artifacts" in heading_texts
    assert "Feature Groups" in heading_texts
    assert "Recommendations" in heading_texts


def test_word_export_contains_feature_details(db_session):
    """Word export should include feature names and details."""
    asmt = _seed_full_assessment(db_session)
    data = generate_word_report(db_session, asmt.id)
    doc = Document(io.BytesIO(data))

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Export Feature Group" in all_text


def test_word_export_contains_recommendation(db_session):
    """Word export should include general recommendations."""
    asmt = _seed_full_assessment(db_session)
    data = generate_word_report(db_session, asmt.id)
    doc = Document(io.BytesIO(data))

    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Upgrade Script Includes" in all_text


def test_word_export_empty_assessment(db_session):
    """Word export of empty assessment should succeed gracefully."""
    asmt = _seed_empty_assessment(db_session)
    data = generate_word_report(db_session, asmt.id)
    doc = Document(io.BytesIO(data))
    assert len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# API endpoint tests
# ---------------------------------------------------------------------------


def test_export_xlsx_endpoint(client, db_session):
    """GET /api/assessments/{id}/export/xlsx should return xlsx content."""
    asmt = _seed_full_assessment(db_session)
    resp = client.get(f"/api/assessments/{asmt.id}/export/xlsx")
    assert resp.status_code == 200
    assert "spreadsheetml" in resp.headers["content-type"]
    assert "attachment" in resp.headers["content-disposition"]
    assert "ASMT0099902" in resp.headers["content-disposition"]

    # Verify it's valid xlsx
    wb = load_workbook(io.BytesIO(resp.content))
    assert "Summary" in wb.sheetnames


def test_export_docx_endpoint(client, db_session):
    """GET /api/assessments/{id}/export/docx should return docx content."""
    asmt = _seed_full_assessment(db_session)
    resp = client.get(f"/api/assessments/{asmt.id}/export/docx")
    assert resp.status_code == 200
    assert "wordprocessingml" in resp.headers["content-type"]
    assert "attachment" in resp.headers["content-disposition"]

    # Verify it's valid docx
    doc = Document(io.BytesIO(resp.content))
    assert len(doc.paragraphs) > 0


def test_export_invalid_format_returns_400(client, db_session):
    """GET /api/assessments/{id}/export/pdf should return 400."""
    asmt = _seed_full_assessment(db_session)
    resp = client.get(f"/api/assessments/{asmt.id}/export/pdf")
    assert resp.status_code == 400


def test_export_nonexistent_assessment_returns_404(client, db_session):
    """Export of nonexistent assessment should return 404."""
    resp = client.get("/api/assessments/999999/export/xlsx")
    assert resp.status_code == 404


def test_export_buttons_in_template(client, db_session):
    """Assessment detail page should contain export buttons."""
    asmt = _seed_full_assessment(db_session)
    resp = client.get(f"/assessments/{asmt.id}")
    assert resp.status_code == 200
    html = resp.text
    assert "Export Excel" in html
    assert "Export Word" in html
    assert f"/api/assessments/{asmt.id}/export/xlsx" in html
    assert f"/api/assessments/{asmt.id}/export/docx" in html
