"""Assessment report export service.

Generates Excel (.xlsx) and Word (.docx) exports of assessment report data.
Reads from the GeneralRecommendation record with category="assessment_report"
and supplements with live database queries for detailed breakdowns.
"""

import io
import json
from datetime import datetime
from typing import Any, Dict, List, Optional

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlmodel import Session, select

from ..models import (
    Assessment,
    Feature,
    FeatureScanResult,
    GeneralRecommendation,
    Instance,
    OriginType,
    ReviewStatus,
    Scan,
    ScanResult,
)


# ---------------------------------------------------------------------------
# Shared data extraction
# ---------------------------------------------------------------------------

def _load_report_data(session: Session, assessment_id: int) -> Dict[str, Any]:
    """Load all data needed for export from the database."""
    assessment = session.get(Assessment, assessment_id)
    if not assessment:
        raise ValueError(f"Assessment not found: {assessment_id}")

    instance = session.get(Instance, assessment.instance_id)

    # Load stored report JSON (if available)
    report_rec = session.exec(
        select(GeneralRecommendation)
        .where(GeneralRecommendation.assessment_id == assessment_id)
        .where(GeneralRecommendation.category == "assessment_report")
    ).first()

    stored_report: Dict[str, Any] = {}
    stored_report_text: Optional[str] = None
    if report_rec and report_rec.description:
        try:
            stored_report = json.loads(report_rec.description)
        except (json.JSONDecodeError, TypeError):
            # If it's rich text (from prompt integration), store as text
            stored_report_text = report_rec.description

    # Load scan results for detailed breakdown
    all_results = session.exec(
        select(ScanResult)
        .join(Scan, ScanResult.scan_id == Scan.id)
        .where(Scan.assessment_id == assessment_id)
    ).all()

    customized_types = [OriginType.modified_ootb, OriginType.net_new_customer]
    customized = [sr for sr in all_results if sr.origin_type in customized_types]

    # Load features
    features = session.exec(
        select(Feature).where(Feature.assessment_id == assessment_id)
    ).all()

    # Load general recommendations (excluding report itself)
    gen_recs = session.exec(
        select(GeneralRecommendation)
        .where(GeneralRecommendation.assessment_id == assessment_id)
        .where(GeneralRecommendation.category != "assessment_report")
    ).all()

    return {
        "assessment": assessment,
        "instance": instance,
        "stored_report": stored_report,
        "stored_report_text": stored_report_text,
        "all_results": all_results,
        "customized": customized,
        "features": features,
        "general_recommendations": gen_recs,
    }


# ---------------------------------------------------------------------------
# Excel export
# ---------------------------------------------------------------------------

# Style constants
_HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
_HEADER_FILL = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
_HEADER_ALIGNMENT = Alignment(horizontal="center", vertical="center", wrap_text=True)
_THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)
_TITLE_FONT = Font(name="Calibri", size=14, bold=True, color="2F5496")
_SUBTITLE_FONT = Font(name="Calibri", size=11, color="404040")


def _style_header_row(ws, row_num: int, col_count: int) -> None:
    """Apply header styling to a row."""
    for col in range(1, col_count + 1):
        cell = ws.cell(row=row_num, column=col)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = _HEADER_ALIGNMENT
        cell.border = _THIN_BORDER


def _auto_width(ws, min_width: int = 12, max_width: int = 50) -> None:
    """Auto-fit column widths based on content."""
    for col_cells in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col_cells[0].column)
        for cell in col_cells:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, min_width), max_width)


def generate_excel_report(session: Session, assessment_id: int) -> bytes:
    """Generate an Excel workbook with multiple sheets for the assessment report.

    Returns the workbook as bytes (xlsx format).
    """
    data = _load_report_data(session, assessment_id)
    assessment = data["assessment"]
    instance = data["instance"]

    wb = Workbook()

    # ── Sheet 1: Summary ──
    ws_summary = wb.active
    ws_summary.title = "Summary"

    ws_summary.cell(row=1, column=1, value="Technical Assessment Report").font = _TITLE_FONT
    ws_summary.cell(row=2, column=1, value=f"Assessment: {assessment.name}").font = _SUBTITLE_FONT
    ws_summary.cell(row=3, column=1, value=f"Number: {assessment.number}").font = _SUBTITLE_FONT
    ws_summary.cell(
        row=4, column=1,
        value=f"Instance: {instance.name if instance else 'N/A'}",
    ).font = _SUBTITLE_FONT
    ws_summary.cell(
        row=5, column=1,
        value=f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
    ).font = _SUBTITLE_FONT

    # Statistics table
    row = 7
    stats_headers = ["Metric", "Value"]
    for ci, h in enumerate(stats_headers, 1):
        ws_summary.cell(row=row, column=ci, value=h)
    _style_header_row(ws_summary, row, len(stats_headers))

    stats_data = [
        ("Total Artifacts", len(data["all_results"])),
        ("Customized Artifacts", len(data["customized"])),
        ("Feature Groups", len(data["features"])),
        ("General Recommendations", len(data["general_recommendations"])),
    ]

    # Add review status counts
    reviewed = sum(1 for sr in data["customized"] if sr.review_status == ReviewStatus.reviewed)
    stats_data.append(("Reviewed Customizations", reviewed))
    stats_data.append(("Pending Review", len(data["customized"]) - reviewed))

    for si, (metric, value) in enumerate(stats_data, row + 1):
        ws_summary.cell(row=si, column=1, value=metric)
        ws_summary.cell(row=si, column=2, value=value)

    _auto_width(ws_summary)

    # ── Sheet 2: Customized Artifacts ──
    ws_artifacts = wb.create_sheet("Customized Artifacts")

    artifact_headers = [
        "Name", "Table", "Origin", "Active", "Review Status",
        "Disposition", "Target Table", "Description",
    ]
    for ci, h in enumerate(artifact_headers, 1):
        ws_artifacts.cell(row=1, column=ci, value=h)
    _style_header_row(ws_artifacts, 1, len(artifact_headers))

    for ri, sr in enumerate(data["customized"], 2):
        ws_artifacts.cell(row=ri, column=1, value=sr.name)
        ws_artifacts.cell(row=ri, column=2, value=sr.table_name)
        ws_artifacts.cell(
            row=ri, column=3,
            value=sr.origin_type.value if sr.origin_type else "",
        )
        ws_artifacts.cell(row=ri, column=4, value="Yes" if sr.is_active else "No")
        ws_artifacts.cell(
            row=ri, column=5,
            value=sr.review_status.value if sr.review_status else "pending",
        )
        ws_artifacts.cell(
            row=ri, column=6,
            value=sr.disposition.value if sr.disposition else "",
        )
        ws_artifacts.cell(row=ri, column=7, value=sr.meta_target_table or "")
        ws_artifacts.cell(row=ri, column=8, value=sr.finding_description or "")

    _auto_width(ws_artifacts)

    # ── Sheet 3: Features ──
    ws_features = wb.create_sheet("Features")

    feature_headers = [
        "Name", "Description", "Disposition", "Member Count",
        "AI Summary", "Recommendation",
    ]
    for ci, h in enumerate(feature_headers, 1):
        ws_features.cell(row=1, column=ci, value=h)
    _style_header_row(ws_features, 1, len(feature_headers))

    for ri, feat in enumerate(data["features"], 2):
        member_count = len(session.exec(
            select(FeatureScanResult).where(FeatureScanResult.feature_id == feat.id)
        ).all())
        ws_features.cell(row=ri, column=1, value=feat.name)
        ws_features.cell(row=ri, column=2, value=feat.description or "")
        ws_features.cell(
            row=ri, column=3,
            value=feat.disposition.value if feat.disposition else "",
        )
        ws_features.cell(row=ri, column=4, value=member_count)
        ws_features.cell(row=ri, column=5, value=feat.ai_summary or "")
        ws_features.cell(row=ri, column=6, value=feat.recommendation or "")

    _auto_width(ws_features)

    # ── Sheet 4: Recommendations ──
    ws_recs = wb.create_sheet("Recommendations")

    rec_headers = ["Title", "Category", "Severity", "Description"]
    for ci, h in enumerate(rec_headers, 1):
        ws_recs.cell(row=1, column=ci, value=h)
    _style_header_row(ws_recs, 1, len(rec_headers))

    for ri, rec in enumerate(data["general_recommendations"], 2):
        ws_recs.cell(row=ri, column=1, value=rec.title)
        ws_recs.cell(row=ri, column=2, value=rec.category or "")
        ws_recs.cell(
            row=ri, column=3,
            value=rec.severity.value if rec.severity else "",
        )
        ws_recs.cell(row=ri, column=4, value=rec.description or "")

    _auto_width(ws_recs)

    # Write to bytes
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------

def generate_word_report(session: Session, assessment_id: int) -> bytes:
    """Generate a Word document for the assessment report.

    Returns the document as bytes (docx format).
    """
    data = _load_report_data(session, assessment_id)
    assessment = data["assessment"]
    instance = data["instance"]

    doc = Document()

    # Title
    title = doc.add_heading("Technical Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Assessment info
    doc.add_paragraph(
        f"Assessment: {assessment.name}\n"
        f"Number: {assessment.number}\n"
        f"Instance: {instance.name if instance else 'N/A'}\n"
        f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    )

    doc.add_paragraph("")  # spacer

    # If we have rich text from prompt integration, include it
    if data["stored_report_text"]:
        doc.add_heading("Report Content", level=1)
        # Split into paragraphs and add
        for para_text in data["stored_report_text"].split("\n\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped.lstrip("# "), level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped.lstrip("# "), level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped.lstrip("# "), level=3)
            elif stripped.startswith("- "):
                for line in stripped.split("\n"):
                    if line.strip().startswith("- "):
                        doc.add_paragraph(line.strip().lstrip("- "), style="List Bullet")
                    else:
                        doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph(stripped)

    # ── Section: Executive Summary ──
    doc.add_heading("Executive Summary", level=1)

    stats_para = doc.add_paragraph()
    stats_para.add_run("Total Artifacts: ").bold = True
    stats_para.add_run(f"{len(data['all_results'])}\n")
    stats_para.add_run("Customized Artifacts: ").bold = True
    stats_para.add_run(f"{len(data['customized'])}\n")
    stats_para.add_run("Feature Groups: ").bold = True
    stats_para.add_run(f"{len(data['features'])}\n")
    stats_para.add_run("Recommendations: ").bold = True
    stats_para.add_run(f"{len(data['general_recommendations'])}")

    reviewed = sum(1 for sr in data["customized"] if sr.review_status == ReviewStatus.reviewed)
    review_para = doc.add_paragraph()
    review_para.add_run("Review Status: ").bold = True
    review_para.add_run(f"{reviewed}/{len(data['customized'])} customizations reviewed")

    # ── Section: Customized Artifacts ──
    doc.add_heading("Customized Artifacts", level=1)

    if data["customized"]:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, header in enumerate(["Name", "Table", "Origin", "Review Status", "Disposition"]):
            hdr[i].text = header
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for sr in data["customized"]:
            row = table.add_row().cells
            row[0].text = sr.name or ""
            row[1].text = sr.table_name or ""
            row[2].text = sr.origin_type.value if sr.origin_type else ""
            row[3].text = sr.review_status.value if sr.review_status else "pending"
            row[4].text = sr.disposition.value if sr.disposition else ""
    else:
        doc.add_paragraph("No customized artifacts found.")

    # ── Section: Feature Groups ──
    doc.add_heading("Feature Groups", level=1)

    if data["features"]:
        for feat in data["features"]:
            member_count = len(session.exec(
                select(FeatureScanResult).where(FeatureScanResult.feature_id == feat.id)
            ).all())
            doc.add_heading(feat.name, level=2)
            doc.add_paragraph(
                f"Disposition: {feat.disposition.value if feat.disposition else 'unassigned'} | "
                f"Members: {member_count}"
            )
            if feat.description:
                doc.add_paragraph(feat.description)
            if feat.ai_summary:
                ai_para = doc.add_paragraph()
                ai_para.add_run("AI Summary: ").bold = True
                ai_para.add_run(feat.ai_summary)
            if feat.recommendation:
                rec_para = doc.add_paragraph()
                rec_para.add_run("Recommendation: ").bold = True
                rec_para.add_run(feat.recommendation)
    else:
        doc.add_paragraph("No feature groups defined.")

    # ── Section: Recommendations ──
    doc.add_heading("Recommendations", level=1)

    if data["general_recommendations"]:
        severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
        sorted_recs = sorted(
            data["general_recommendations"],
            key=lambda r: severity_order.get(r.severity.value if r.severity else "info", 5),
        )
        for rec in sorted_recs:
            sev = f"[{rec.severity.value}]" if rec.severity else "[unrated]"
            doc.add_heading(f"{sev} {rec.title}", level=2)
            if rec.description:
                doc.add_paragraph(rec.description)
    else:
        doc.add_paragraph("No recommendations recorded.")

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
