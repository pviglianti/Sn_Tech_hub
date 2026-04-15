"""Server-side report generation for an Assessment.

Produces:
  - {DATA_DIR}/reports/{assessment_id}/assessment_{number}_report_{timestamp}.xlsx
  - {DATA_DIR}/reports/{assessment_id}/assessment_{number}_report_{timestamp}.docx

Each generation creates new AssessmentReport rows so history is preserved.
"""

from __future__ import annotations

import hashlib
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from sqlmodel import Session, select

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt

from ..database import DATA_DIR
from ..models import (
    Assessment,
    AssessmentReport,
    Customization,
    Feature,
    FeatureScanResult,
    GlobalApp,
    ScanResult,
    Scan,
)


logger = logging.getLogger(__name__)


def _reports_dir(assessment_id: int) -> Path:
    out = DATA_DIR / "reports" / str(assessment_id)
    out.mkdir(parents=True, exist_ok=True)
    return out


def _load_json_list(raw: Optional[str]) -> List[Any]:
    if not raw:
        return []
    try:
        v = json.loads(raw)
        return v if isinstance(v, list) else []
    except json.JSONDecodeError:
        return []


def _gather_report_data(session: Session, assessment_id: int) -> Dict[str, Any]:
    assessment = session.get(Assessment, assessment_id)
    if assessment is None:
        raise ValueError(f"Assessment {assessment_id} not found")

    target_app: Optional[GlobalApp] = None
    if assessment.target_app_id:
        target_app = session.get(GlobalApp, assessment.target_app_id)

    # The "Customizations" related list on the assessment is the source of truth
    # for what's reportable. Customization is a child table of ScanResult containing
    # only customized rows (origin_type in {modified_ootb, net_new_customer}).
    # Out-of-scope rows are filtered OUT of the customizations tab in the UI; here
    # we keep them so we can still render the "Out of Scope" tab.
    scan_ids = [s.id for s in assessment.scans] if assessment.scans else []
    customized: List[Customization] = []
    if scan_ids:
        customized = list(session.exec(
            select(Customization).where(Customization.scan_id.in_(scan_ids))
        ).all())
    # Keep this name for legacy reads below (e.g. anywhere the report iterates rows)
    results: List[Customization] = customized

    features = list(session.exec(
        select(Feature).where(Feature.assessment_id == assessment_id)
    ).all())
    features_by_id = {f.id: f for f in features}

    # Build {scan_result_id: feature} map via the M2M join (primary membership preferred)
    feature_id_set = {f.id for f in features if f.id is not None}
    membership_rows = []
    if feature_id_set:
        membership_rows = list(session.exec(
            select(FeatureScanResult).where(FeatureScanResult.feature_id.in_(feature_id_set))
        ).all())
    result_to_feature: Dict[int, Feature] = {}
    for m in membership_rows:
        # Prefer primary membership; otherwise first-seen wins
        existing = result_to_feature.get(m.scan_result_id)
        if existing is None or m.is_primary:
            feat = features_by_id.get(m.feature_id)
            if feat is not None:
                result_to_feature[m.scan_result_id] = feat

    # Bucket strictly by the two scope checkboxes the triage step sets:
    #   - is_out_of_scope=True → Out of Scope tab
    #   - is_adjacent=True     → reportable as adjacent (still in-scope)
    #   - neither              → reportable as in_scope
    in_scope: List[Customization] = []
    adjacent: List[Customization] = []
    out_of_scope: List[Customization] = []
    for r in customized:
        if getattr(r, "is_out_of_scope", False):
            out_of_scope.append(r)
        elif getattr(r, "is_adjacent", False):
            adjacent.append(r)
        else:
            in_scope.append(r)

    in_scope_reportable = in_scope + adjacent  # what goes on the In-Scope tab

    return {
        "assessment": assessment,
        "target_app": target_app,
        "core_tables": _load_json_list(target_app.core_tables_json) if target_app else [],
        "explicit_target_tables": _load_json_list(assessment.target_tables_json),
        "file_classes": _load_json_list(assessment.app_file_classes_json),
        "results": results,
        "in_scope": in_scope,
        "adjacent": adjacent,
        "out_of_scope": out_of_scope,
        "customized": customized,
        "features": features,
        "features_by_id": features_by_id,
        "result_to_feature": result_to_feature,
        "in_scope_reportable": in_scope_reportable,
    }


# ────────────────────────────────────────────────────────────────────────
# Excel
# ────────────────────────────────────────────────────────────────────────

_HEADER_FONT = Font(bold=True, color="FFFFFF")
_HEADER_FILL = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")


def _autosize(ws, headers: List[str]) -> None:
    for i, h in enumerate(headers, start=1):
        col = get_column_letter(i)
        max_len = len(str(h))
        for row in ws.iter_rows(min_row=2, min_col=i, max_col=i, values_only=True):
            v = row[0]
            if v is None:
                continue
            l = len(str(v))
            if l > max_len:
                max_len = l
        ws.column_dimensions[col].width = min(max(max_len + 2, 10), 80)


def _write_header(ws, headers: List[str]) -> None:
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=i, value=h)
        cell.font = _HEADER_FONT
        cell.fill = _HEADER_FILL
        cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def _build_xlsx(data: Dict[str, Any]) -> Workbook:
    wb = Workbook()
    a: Assessment = data["assessment"]
    target_app = data["target_app"]

    # ─── Tab 1: Executive Summary ──────────────────────────
    ws = wb.active
    ws.title = "Executive Summary"
    title_font = Font(bold=True, size=18)
    section_font = Font(bold=True, size=12, color="1F4E78")
    label_font = Font(bold=True)

    title_cell = ws.cell(row=1, column=1, value=f"Technical Assessment Report: {a.name}")
    title_cell.font = title_font
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=4)
    ws.cell(row=2, column=1, value=f"Assessment {a.number}  •  Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    # Scope section
    r = 4
    ws.cell(row=r, column=1, value="SCOPE").font = section_font
    r += 1
    scope_pairs = [
        ("Target Application", target_app.label if target_app else "—"),
        ("Assessment Type", str(a.assessment_type.value if hasattr(a.assessment_type, "value") else a.assessment_type)),
        ("Core Tables", ", ".join(data["core_tables"]) or "—"),
        ("Parent Table", (target_app.parent_table if target_app else None) or "—"),
        ("Scope Filter", a.scope_filter),
        ("Pipeline Stage", str(a.pipeline_stage.value if hasattr(a.pipeline_stage, "value") else a.pipeline_stage)),
    ]
    for k, v in scope_pairs:
        ws.cell(row=r, column=1, value=k).font = label_font
        ws.cell(row=r, column=2, value=v)
        r += 1

    # Counts section
    r += 1
    ws.cell(row=r, column=1, value="ARTIFACT COUNTS").font = section_font
    r += 1
    customized = data["customized"]
    in_scope_plus_adj = data["in_scope_reportable"]
    count_pairs = [
        ("Total Scanned", len(data["results"])),
        ("Customized (in scope + adjacent + ootb-modified)", len(customized)),
        ("In Scope", len(data["in_scope"])),
        ("Adjacent (parent/related table)", len(data["adjacent"])),
        ("Out of Scope", len(data["out_of_scope"])),
        ("Features Identified", len(data["features"])),
    ]
    for k, v in count_pairs:
        ws.cell(row=r, column=1, value=k).font = label_font
        ws.cell(row=r, column=2, value=v)
        r += 1

    # Top features section
    r += 1
    ws.cell(row=r, column=1, value="TOP FEATURES BY ARTIFACT COUNT").font = section_font
    r += 1
    feat_counts: Dict[int, int] = {}
    feat_types: Dict[int, set] = {}
    for sr in customized:
        feat = data["result_to_feature"].get(getattr(sr, "scan_result_id", sr.id))
        if feat and feat.id is not None:
            feat_counts[feat.id] = feat_counts.get(feat.id, 0) + 1
            feat_types.setdefault(feat.id, set()).add(getattr(sr, "sys_class_name", None) or "?")
    top = sorted(data["features"], key=lambda f: -feat_counts.get(f.id, 0))[:10]
    ws.cell(row=r, column=1, value="Feature").font = label_font
    ws.cell(row=r, column=2, value="Count").font = label_font
    ws.cell(row=r, column=3, value="Risk").font = label_font
    r += 1
    for f in top:
        ws.cell(row=r, column=1, value=f.name)
        ws.cell(row=r, column=2, value=feat_counts.get(f.id, 0))
        ws.cell(row=r, column=3, value=getattr(f, "change_risk_level", None) or "—")
        r += 1

    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 18

    # ─── Tab 2: Feature Inventory ──────────────────────────
    ws = wb.create_sheet("Feature Inventory")
    headers = ["Feature Name", "Description", "Artifact Count", "Types", "Risk Level", "Key Risks", "Recommendation", "AI Summary"]
    _write_header(ws, headers)
    feats_sorted = sorted(data["features"], key=lambda f: -feat_counts.get(f.id, 0))
    for row_idx, f in enumerate(feats_sorted, start=2):
        types = sorted(feat_types.get(f.id, set()))
        ws.cell(row=row_idx, column=1, value=f.name)
        ws.cell(row=row_idx, column=2, value=f.description)
        ws.cell(row=row_idx, column=3, value=feat_counts.get(f.id, 0))
        ws.cell(row=row_idx, column=4, value=", ".join(types))
        ws.cell(row=row_idx, column=5, value=getattr(f, "change_risk_level", None))
        ws.cell(row=row_idx, column=6, value=getattr(f, "key_risks", None) or getattr(f, "ai_risks", None))
        ws.cell(row=row_idx, column=7, value=getattr(f, "recommendation", None))
        ws.cell(row=row_idx, column=8, value=getattr(f, "ai_summary", None))
    _autosize(ws, headers)

    # ─── Tab 3: In-Scope Customizations ───────────────────
    ws = wb.create_sheet("In-Scope Customizations")
    headers = ["ID", "Name", "Table", "Class", "Origin Type", "Scope", "Feature Name", "Observations", "Recommendation"]
    _write_header(ws, headers)
    for row_idx, sr in enumerate(in_scope_plus_adj, start=2):
        scope = "adjacent" if getattr(sr, "is_adjacent", False) else "in_scope"
        feat = data["result_to_feature"].get(getattr(sr, "scan_result_id", sr.id))
        ws.cell(row=row_idx, column=1, value=sr.id)
        ws.cell(row=row_idx, column=2, value=getattr(sr, "name", None))
        ws.cell(row=row_idx, column=3, value=getattr(sr, "table_name", None))
        ws.cell(row=row_idx, column=4, value=getattr(sr, "sys_class_name", None))
        ws.cell(row=row_idx, column=5, value=str(getattr(sr, "origin_type", None) or ""))
        ws.cell(row=row_idx, column=6, value=scope)
        ws.cell(row=row_idx, column=7, value=feat.name if feat else "")
        ws.cell(row=row_idx, column=8, value=getattr(sr, "observations", None))
        ws.cell(row=row_idx, column=9, value=getattr(sr, "recommendation", None))
    _autosize(ws, headers)

    # ─── Tab 4: Out of Scope ──────────────────────────────
    ws = wb.create_sheet("Out of Scope")
    headers = ["ID", "Name", "Table", "Class", "Origin Type", "Observations"]
    _write_header(ws, headers)
    for row_idx, sr in enumerate(data["out_of_scope"], start=2):
        ws.cell(row=row_idx, column=1, value=sr.id)
        ws.cell(row=row_idx, column=2, value=getattr(sr, "name", None))
        ws.cell(row=row_idx, column=3, value=getattr(sr, "table_name", None))
        ws.cell(row=row_idx, column=4, value=getattr(sr, "sys_class_name", None))
        ws.cell(row=row_idx, column=5, value=str(getattr(sr, "origin_type", None) or ""))
        ws.cell(row=row_idx, column=6, value=getattr(sr, "observations", None))
    _autosize(ws, headers)

    # ─── Tab 5: Risk Matrix ──────────────────────────────
    ws = wb.create_sheet("Risk Matrix")
    headers = ["Risk Category", "Count", "Severity", "Affected Features", "Example Artifact"]
    _write_header(ws, headers)
    # Bucket features by risk level
    risk_buckets: Dict[str, List[Feature]] = {}
    for f in data["features"]:
        lvl = (getattr(f, "change_risk_level", None) or "unspecified").lower()
        risk_buckets.setdefault(lvl, []).append(f)
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "unspecified": 4}
    risk_levels = sorted(risk_buckets.keys(), key=lambda k: severity_order.get(k, 99))
    for row_idx, lvl in enumerate(risk_levels, start=2):
        feats = risk_buckets[lvl]
        affected_count = sum(feat_counts.get(f.id, 0) for f in feats)
        example = ""
        for f in feats:
            for sr in customized:
                if data["result_to_feature"].get(getattr(sr, "scan_result_id", sr.id)) is f:
                    example = f"{getattr(sr, 'name', '')} ({getattr(sr, 'sys_class_name', '')})"
                    break
            if example:
                break
        ws.cell(row=row_idx, column=1, value=f"Change Risk: {lvl}")
        ws.cell(row=row_idx, column=2, value=affected_count)
        ws.cell(row=row_idx, column=3, value=lvl)
        ws.cell(row=row_idx, column=4, value=", ".join(f.name for f in feats[:5]))
        ws.cell(row=row_idx, column=5, value=example)
    _autosize(ws, headers)

    return wb


# ────────────────────────────────────────────────────────────────────────
# Word
# ────────────────────────────────────────────────────────────────────────

def _build_docx(data: Dict[str, Any]) -> Document:
    import re
    doc = Document()
    a: Assessment = data["assessment"]
    target_app = data["target_app"]
    customized = data["customized"]
    in_scope_reportable = data["in_scope_reportable"]

    # ── Title ─────────────────────────────────────────────
    doc.add_paragraph("Technical Assessment Report")
    doc.add_paragraph(a.name or "Assessment")
    doc.add_paragraph(f"Prepared: {datetime.utcnow().strftime('%B %d, %Y')}")
    if target_app:
        doc.add_paragraph(f"ServiceNow {target_app.label}")

    # ── Executive Summary narrative ──────────────────────
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"This assessment analyzed {len(customized)} customized artifacts across "
        f"{target_app.label if target_app else 'the target application'}. "
        f"Of these, {len(in_scope_reportable)} customizations "
        f"({len(data['in_scope'])} direct, {len(data['adjacent'])} adjacent) "
        f"were classified as in-scope for review, with "
        f"{len(data['out_of_scope'])} marked out-of-scope. "
        f"Reportable customizations were grouped into {len(data['features'])} "
        f"business features."
    )

    # ── Key Metrics ──────────────────────────────────────
    doc.add_heading("Key Metrics", level=2)
    metrics = [
        ("Customizations", len(customized)),
        ("In Scope (direct)", len(data["in_scope"])),
        ("Adjacent (parent/related)", len(data["adjacent"])),
        ("Out of Scope", len(data["out_of_scope"])),
        ("Features Identified", len(data["features"])),
    ]
    for label, val in metrics:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(str(val))

    # ── Recommendation Breakdown (parse disposition keywords from text) ──
    def _classify_recommendation(text: Optional[str]) -> str:
        if not text:
            return "none"
        t = text.lower()
        # Word-boundary matches; order matters for ambiguous cases
        if re.search(r"\bretire\b|\bremove\b|\bdelete\b", t):
            return "retire"
        if re.search(r"\breplace (with )?ootb\b|\breplace with the platform\b|\buse ootb\b", t):
            return "replace"
        if re.search(r"\brefactor\b|\bclean up\b|\brework\b|\bmigrate to\b", t):
            return "refactor"
        if re.search(r"\bkeep (as[- ]is|unchanged)?\b|\bfollows best practices\b|\bwell[- ]implemented\b", t):
            return "keep"
        return "review"

    disp_counts = {"keep": 0, "refactor": 0, "replace": 0, "retire": 0, "review": 0, "none": 0}
    for sr in in_scope_reportable:
        disp_counts[_classify_recommendation(getattr(sr, "recommendation", None))] += 1

    total_in = len(in_scope_reportable) or 1
    doc.add_heading("Recommendation Breakdown", level=2)
    doc.add_paragraph(
        f"Of the {len(in_scope_reportable)} in-scope artifacts: "
        f"{disp_counts['keep']} ({100 * disp_counts['keep'] // total_in}%) follow best practices "
        f"and should be kept as-is, "
        f"{disp_counts['refactor']} ({100 * disp_counts['refactor'] // total_in}%) need refactoring, "
        f"{disp_counts['replace']} require replacement with OOTB, and "
        f"{disp_counts['retire']} should be retired."
    )

    # ── Top Findings (regex patterns across recommendations + ai_observations) ──
    doc.add_heading("Top 5 Findings", level=2)
    FINDINGS = [
        ("setWorkflow(false) usage",
         r"setworkflow\s*\(\s*false\s*\)",
         "disables audit trails and notifications"),
        ("Hardcoded sys_ids",
         r"hardcoded sys[-_ ]?id|hard[- ]coded reference|[0-9a-f]{32}",
         "creates fragile dependencies"),
        ("Overly permissive ACLs",
         r"acl.*no role|acls? (lack|without)|no role requirement",
         "grants unrestricted access"),
        ("current.update() in business rules",
         r"current\.update\(\)",
         "can cause recursion / double processing"),
        ("GlideRecord in loops",
         r"gliderecord (query )?in a? loop|n\+1",
         "performance concern"),
        ("Missing error handling",
         r"no (try|error) handling|missing try[/ ]?catch",
         "fragile under failure conditions"),
    ]
    top_findings = []
    for title, pattern, why in FINDINGS:
        count = 0
        for sr in in_scope_reportable:
            text_blob = " ".join(
                str(getattr(sr, fld, None) or "")
                for fld in ("observations", "recommendation")
            ).lower()
            if re.search(pattern, text_blob):
                count += 1
        if count:
            top_findings.append((count, title, why))
    top_findings.sort(key=lambda x: -x[0])
    for i, (count, title, why) in enumerate(top_findings[:5], start=1):
        doc.add_paragraph(f"{i}. {title}: {count} artifacts — {why}.")
    if not top_findings:
        doc.add_paragraph("(No patterned findings extracted from recommendation text.)")

    # ── Feature-by-Feature Analysis ───────────────────────
    doc.add_page_break()
    doc.add_heading("Feature-by-Feature Analysis", level=1)
    feat_counts: Dict[int, int] = {}
    feat_types: Dict[int, Dict[str, int]] = {}
    for r in customized:
        feat = data["result_to_feature"].get(getattr(r, "scan_result_id", r.id))
        if feat and feat.id is not None:
            feat_counts[feat.id] = feat_counts.get(feat.id, 0) + 1
            feat_types.setdefault(feat.id, {}).setdefault(getattr(r, "sys_class_name", None) or "?", 0)
            feat_types[feat.id][getattr(r, "sys_class_name", None) or "?"] += 1

    for f in sorted(data["features"], key=lambda f: -feat_counts.get(f.id, 0)):
        doc.add_heading(f.name or f"Feature {f.id}", level=2)
        types_map = feat_types.get(f.id, {})
        types_str = ", ".join(f"{k} ({v})" for k, v in sorted(types_map.items(), key=lambda kv: -kv[1]))
        meta = doc.add_paragraph()
        meta.add_run(
            f"Artifacts: {feat_counts.get(f.id, 0)} | "
            f"Risk: {getattr(f, 'change_risk_level', None) or '—'} | "
            f"Types: {types_str or '—'}"
        )
        if getattr(f, "description", None):
            doc.add_paragraph(f.description)
        if getattr(f, "recommendation", None):
            p = doc.add_paragraph()
            r = p.add_run("Recommendation: ")
            r.bold = True
            p.add_run(f.recommendation)
        if getattr(f, "ai_summary", None):
            p = doc.add_paragraph()
            r = p.add_run("AI Summary: ")
            r.bold = True
            p.add_run(f.ai_summary)

    # ── Appendix: Counts by Table ─────────────────────────
    doc.add_page_break()
    doc.add_heading("Appendix — Counts by Table", level=1)
    by_table: Dict[str, int] = {}
    for r in customized:
        t = getattr(r, "table_name", None) or "(no table)"
        by_table[t] = by_table.get(t, 0) + 1
    for t, n in sorted(by_table.items(), key=lambda kv: -kv[1]):
        doc.add_paragraph(f"{t}: {n}", style="List Bullet")

    return doc


# ────────────────────────────────────────────────────────────────────────
# Public API
# ────────────────────────────────────────────────────────────────────────

def generate_reports(
    session: Session,
    assessment_id: int,
    formats: Optional[List[str]] = None,
    generated_by: Optional[str] = None,
) -> List[AssessmentReport]:
    """Generate one or more report formats for an assessment.

    Returns the persisted AssessmentReport rows.
    """
    formats = [f.lower() for f in (formats or ["xlsx", "docx"]) if f]
    if not formats:
        raise ValueError("At least one format must be specified (xlsx, docx)")
    for f in formats:
        if f not in ("xlsx", "docx"):
            raise ValueError(f"Unsupported format: {f}")

    data = _gather_report_data(session, assessment_id)
    a: Assessment = data["assessment"]
    out_dir = _reports_dir(assessment_id)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    base = f"assessment_{a.number}_{timestamp}"

    created: List[AssessmentReport] = []

    if "xlsx" in formats:
        wb = _build_xlsx(data)
        path = out_dir / f"{base}.xlsx"
        wb.save(path)
        created.append(_persist_report(session, assessment_id, path, "xlsx", generated_by))

    if "docx" in formats:
        doc = _build_docx(data)
        path = out_dir / f"{base}.docx"
        doc.save(path)
        created.append(_persist_report(session, assessment_id, path, "docx", generated_by))

    session.commit()
    for r in created:
        session.refresh(r)
    return created


def _persist_report(
    session: Session,
    assessment_id: int,
    path: Path,
    fmt: str,
    generated_by: Optional[str],
) -> AssessmentReport:
    file_bytes = path.read_bytes()
    sha = hashlib.sha256(file_bytes).hexdigest()
    row = AssessmentReport(
        assessment_id=assessment_id,
        filename=path.name,
        format=fmt,
        file_path=str(path),
        file_size=len(file_bytes),
        sha256=sha,
        generated_by=generated_by,
    )
    session.add(row)
    session.flush()
    return row


def list_reports(session: Session, assessment_id: int) -> List[AssessmentReport]:
    return list(session.exec(
        select(AssessmentReport)
        .where(AssessmentReport.assessment_id == assessment_id)
        .order_by(AssessmentReport.generated_at.desc())
    ).all())


def get_report(session: Session, report_id: int) -> Optional[AssessmentReport]:
    return session.get(AssessmentReport, report_id)
